# -*- coding: utf-8 -*-
# ==========================================================================================
# App: Encuesta Comercio → XLSForm para ArcGIS Survey123 (versión extendida)
# - Constructor completo (agregar/editar/ordenar/borrar)
# - Condicionales (relevant) + finalizar temprano (oculta lo siguiente)
# - Listas en cascada (choice_filter) Cantón→Distrito→Barrio [CATÁLOGO MANUAL POR LOTES]
# - Exportar/Importar proyecto (JSON)
# - Exportar a XLSForm (survey/choices/settings)
# - PÁGINAS reales (style="pages"): P1..P6
# - Portada con logo (media::image) e introducción
# - Word/PDF: imprimen P1 y el resto de páginas
# ==========================================================================================

import re
import json
from io import BytesIO
from datetime import datetime
from typing import List, Dict, Tuple

import streamlit as st
import pandas as pd

# ------------------------------------------------------------------------------------------
# Configuración de la app
# ------------------------------------------------------------------------------------------
st.set_page_config(page_title="Encuesta Comercio → XLSForm (Survey123)", layout="wide")
st.title("🏪 Encuesta Comercio → XLSForm para ArcGIS Survey123")

st.markdown("""
Crea tu cuestionario y **exporta un XLSForm** listo para **ArcGIS Survey123**.

Incluye:
- Tipos: **text**, **integer/decimal**, **date**, **time**, **geopoint**, **select_one**, **select_multiple**.
- **Constructor completo** (agregar, editar, ordenar, borrar) con condicionales.
- **Listas en cascada** **Cantón→Distrito→Barrio** (**catálogo manual por lotes**).
- **Páginas** con navegación **Siguiente/Anterior** (`settings.style = pages`).
- **Portada** con **logo** (`media::image`) e **introducción**.
""")

# ------------------------------------------------------------------------------------------
# Helpers
# ------------------------------------------------------------------------------------------
TIPOS = [
    "Texto (corto)",
    "Párrafo (texto largo)",
    "Número",
    "Selección única",
    "Selección múltiple",
    "Fecha",
    "Hora",
    "GPS (ubicación)"
]

def _rerun():
    if hasattr(st, "rerun"):
        st.rerun()
    else:
        st.experimental_rerun()

def slugify_name(texto: str) -> str:
    if not texto:
        return "campo"
    t = texto.lower()
    t = re.sub(r"[áàäâ]", "a", t); t = re.sub(r"[éèëê]", "e", t)
    t = re.sub(r"[íìïî]", "i", t); t = re.sub(r"[óòöô]", "o", t)
    t = re.sub(r"[úùüû]", "u", t); t = re.sub(r"ñ", "n", t)
    t = re.sub(r"[^a-z0-9]+", "_", t).strip("_")
    return t or "campo"

def asegurar_nombre_unico(base: str, usados: set) -> str:
    if base not in usados: return base
    i = 2
    while f"{base}_{i}" in usados:
        i += 1
    return f"{base}_{i}"

def map_tipo_to_xlsform(tipo_ui: str, name: str):
    if tipo_ui == "Texto (corto)": return ("text", None, None)
    if tipo_ui == "Párrafo (texto largo)": return ("text", "multiline", None)
    if tipo_ui == "Número": return ("integer", None, None)
    if tipo_ui == "Selección única": return (f"select_one list_{name}", None, f"list_{name}")
    if tipo_ui == "Selección múltiple": return (f"select_multiple list_{name}", None, f"list_{name}")
    if tipo_ui == "Fecha": return ("date", None, None)
    if tipo_ui == "Hora": return ("time", None, None)
    if tipo_ui == "GPS (ubicación)": return ("geopoint", None, None)
    return ("text", None, None)

def xlsform_or_expr(conds):
    if not conds: return None
    if len(conds) == 1: return conds[0]
    return "(" + " or ".join(conds) + ")"

def xlsform_not(expr):
    if not expr: return None
    return f"not({expr})"

def build_relevant_expr(rules_for_target: List[Dict]):
    or_parts = []
    for r in rules_for_target:
        src = r["src"]; op = r.get("op", "="); vals = r.get("values", [])
        if not vals: continue
        if op == "=":
            segs = [f"${{{src}}}='{v}'" for v in vals]
        elif op == "selected":
            segs = [f"selected(${{{src}}}, '{v}')" for v in vals]
        elif op == "!=":
            segs = [f"${{{src}}}!='{v}'" for v in vals]
        else:
            segs = [f"${{{src}}}='{v}'" for v in vals]
        or_parts.append(xlsform_or_expr(segs))
    return xlsform_or_expr(or_parts)
# ------------------------------------------------------------------------------------------
# Catálogo manual por lotes: Cantón → Distrito → (Barrios…)
# ------------------------------------------------------------------------------------------
if "choices_ext_rows" not in st.session_state:
    st.session_state.choices_ext_rows = []  # filas para hoja choices
if "choices_extra_cols" not in st.session_state:
    st.session_state.choices_extra_cols = set()

def _append_choice_unique(row: Dict):
    """Inserta fila en choices evitando duplicados por (list_name,name)."""
    key = (row.get("list_name"), row.get("name"))
    exists = any((r.get("list_name"), r.get("name")) == key for r in st.session_state.choices_ext_rows)
    if not exists:
        st.session_state.choices_ext_rows.append(row)

st.markdown("### 📚 Catálogo Cantón → Distrito → Barrios (por lotes)")
with st.expander("Agrega un lote (un Cantón, un Distrito y varios Barrios)", expanded=True):
    col_c1, col_c2 = st.columns(2)
    canton_txt = col_c1.text_input("Cantón (una vez)", value="")
    distrito_txt = col_c2.text_input("Distrito (una vez)", value="")
    barrios_txt = st.text_area("Barrios del distrito (uno por línea)", value="", height=130)

    col_b1, col_b2, col_b3 = st.columns([1,1,2])
    add_lote = col_b1.button("Agregar lote", type="primary", use_container_width=True)
    clear_all = col_b2.button("Limpiar catálogo", use_container_width=True)

    if clear_all:
        st.session_state.choices_ext_rows = []
        st.success("Catálogo limpiado.")

    if add_lote:
        c = canton_txt.strip(); d = distrito_txt.strip()
        barrios = [b.strip() for b in barrios_txt.splitlines() if b.strip()]
        if not c or not d or not barrios:
            st.error("Debes indicar Cantón, Distrito y al menos un Barrio.")
        else:
            slug_c = slugify_name(c); slug_d = slugify_name(d)
            st.session_state.choices_extra_cols.update({"canton_key","distrito_key","any"})

            # Placeholders (una sola vez por lista)
            _append_choice_unique({"list_name":"list_canton", "name":"__pick_canton__", "label":"— escoja un cantón —"})
            _append_choice_unique({"list_name":"list_distrito","name":"__pick_distrito__","label":"— escoja un cantón —","any":"1"})
            _append_choice_unique({"list_name":"list_barrio", "name":"__pick_barrio__", "label":"— escoja un distrito —","any":"1"})

            # Cantón (evita duplicados automáticamente)
            _append_choice_unique({"list_name":"list_canton","name":slug_c,"label":c})

            # Distrito
            _append_choice_unique({"list_name":"list_distrito","name":slug_d,"label":d,"canton_key":slug_c})

            # Barrios
            usados_b = set()
            for b in barrios:
                slug_b = asegurar_nombre_unico(slugify_name(b), usados_b); usados_b.add(slug_b)
                _append_choice_unique({"list_name":"list_barrio","name":slug_b,"label":b,"distrito_key":slug_d})

            st.success(f"Lote agregado: {c} → {d} → {len(barrios)} barrios.")

# Vista previa de catálogo
if st.session_state.choices_ext_rows:
    st.dataframe(pd.DataFrame(st.session_state.choices_ext_rows),
                 use_container_width=True, hide_index=True, height=240)

# ------------------------------------------------------------------------------------------
# Cabecera: Logo + Delegación
# ------------------------------------------------------------------------------------------
DEFAULT_LOGO_PATH = "001.png"

col_logo, col_txt = st.columns([1, 3], vertical_alignment="center")
with col_logo:
    up_logo = st.file_uploader("Logo (PNG/JPG)", type=["png","jpg","jpeg"])
    if up_logo:
        st.image(up_logo, caption="Logo cargado", use_container_width=True)
        st.session_state["_logo_bytes"] = up_logo.getvalue()
        st.session_state["_logo_name"] = up_logo.name
    else:
        try:
            st.image(DEFAULT_LOGO_PATH, caption="Logo (001.png)", use_container_width=True)
            st.session_state["_logo_bytes"] = None
            st.session_state["_logo_name"] = "001.png"
        except Exception:
            st.warning("Sube un logo para incluirlo en el XLSForm.")
            st.session_state["_logo_bytes"] = None
            st.session_state["_logo_name"] = "logo.png"

with col_txt:
    st.markdown("<div style='height:8px'></div>", unsafe_allow_html=True)
    delegacion = st.text_input("Nombre del lugar / Delegación", value="San Carlos Oeste")
    logo_media_name = st.text_input("Nombre de archivo para `media::image`",
                                    value=st.session_state.get("_logo_name","001.png"),
                                    help="Debe coincidir con el archivo en `media/` de Survey123 Connect.")
    titulo_compuesto = (f"Encuesta Comercio – {delegacion.strip()}" if delegacion.strip() else "Encuesta Comercio")
    st.markdown(f"<h5 style='text-align:center;margin:4px 0'>📋 {titulo_compuesto}</h5>", unsafe_allow_html=True)

# ------------------------------------------------------------------------------------------
# Estado
# ------------------------------------------------------------------------------------------
if "preguntas" not in st.session_state: st.session_state.preguntas = []
if "reglas_visibilidad" not in st.session_state: st.session_state.reglas_visibilidad = []
if "reglas_finalizar" not in st.session_state: st.session_state.reglas_finalizar = []
# ------------------------------------------------------------------------------------------
# Intro (Página 1) — COMERCIO
# ------------------------------------------------------------------------------------------
INTRO_COMERCIO = (
    "Con el objetivo de fortalecer la seguridad en nuestro entorno comercial, nos enfocamos "
    "en abordar las principales preocupaciones de seguridad. Es crucial colaborar estrechamente, "
    "no solo con el gobierno local y otras instituciones, sino también con la comunidad empresarial. "
    "Juntos, podemos implementar medidas efectivas para reducir delitos y minimizar riesgos que "
    "impactan directamente en la seguridad de nuestra gente.\n\n"
    "Cabe destacar que la información que nos suministras es completamente confidencial y se empleará "
    "exclusivamente con el propósito de mejorar la seguridad en nuestra área comercial."
)

# ------------------------------------------------------------------------------------------
# Precarga de preguntas (SEED) — Encuesta Comercio (6 páginas)
# Mantiene el catálogo en cascada Cantón→Distrito→Barrio SIN generar opciones aquí.
# ------------------------------------------------------------------------------------------
if "seed_cargado" not in st.session_state or st.session_state.get("perfil_seed") != "comercio":
    v_mas  = slugify_name("Más Seguro")
    v_igual = slugify_name("Igual")
    v_menos = slugify_name("Menos Seguro")

    seed = [
        # ===================== PÁGINA 2: Datos demográficos =====================
        {"tipo_ui":"Selección única","label":"Cantón","name":"canton","required":True,
         "opciones":[], "appearance":None, "choice_filter":None, "relevant":None},

        {"tipo_ui":"Selección única","label":"Distrito","name":"distrito","required":True,
         "opciones":[], "appearance":None, "choice_filter":"canton_key=${canton} or any='1'", "relevant":None},

        {"tipo_ui":"Selección única","label":"Barrio","name":"barrio","required":True,
         "opciones":[], "appearance":None, "choice_filter":"distrito_key=${distrito} or any='1'", "relevant":None},

        {"tipo_ui":"Número","label":"Edad","name":"edad","required":True,
         "opciones":[], "appearance":None, "choice_filter":None, "relevant":None},

        {"tipo_ui":"Selección única","label":"Género","name":"genero","required":True,
         "opciones":["Masculino","Femenino","LGBTQ+"], "appearance":None, "choice_filter":None, "relevant":None},

        {"tipo_ui":"Selección única","label":"Escolaridad","name":"escolaridad","required":True,
         "opciones":["Ninguna","Primaria","Primaria incompleta","Secundaria completa","Secundaria incompleta",
                     "Universitaria","Universitaria incompleta","Técnico"],
         "appearance":None, "choice_filter":None, "relevant":None},

        {"tipo_ui":"Selección única","label":"Tipo de local comercial","name":"tipo_local","required":True,
         "opciones":["Supermercado","Pulpería / Licorera","Restaurante / Soda","Bar","Tienda de artículos",
                     "Gasolineras","Servicios estético","Puesto de lotería","Educación Técnica","Servicio Técnico",
                     "Universidad","Entidad Finaciera","Venta de Servicios","Otros"],
         "appearance":None,"choice_filter":None,"relevant":None},

        # ===================== PÁGINA 3: Sentimiento de inseguridad comercial =====================
        {"tipo_ui":"Selección única","label":"¿Se siente seguro en su local comercial?","name":"seguro_local","required":True,
         "opciones":["Sí","No"], "appearance":None, "choice_filter":None, "relevant":None},

        {"tipo_ui":"Párrafo (texto largo)","label":"Indique por qué se siente inseguro.","name":"motivo_inseguro","required":True,
         "opciones":[], "appearance":None, "choice_filter":None,
         "relevant": f"${{seguro_local}}='{slugify_name('No')}'"},

        {"tipo_ui":"Selección única","label":"¿Cómo se siente en cuanto a la seguridad en su zona comercial este año en comparación con el año anterior?","name":"comparativa_seguridad","required":True,
         "opciones":["Más Seguro","Igual","Menos Seguro"], "appearance":None, "choice_filter":None, "relevant":None},

        {"tipo_ui":"Párrafo (texto largo)","label":"Indique por qué.","name":"motivo_comparativa","required":True,
         "opciones":[], "appearance":None, "choice_filter":None,
         "relevant": "(" + " or ".join([
            f"${{comparativa_seguridad}}='{v_mas}'",
            f"${{comparativa_seguridad}}='{v_igual}'",
            f"${{comparativa_seguridad}}='{v_menos}'"
         ]) + ")"},

        # ===================== PÁGINA 4: Incidencia relacionada a delitos (NO obligatorias) =====================
        {"tipo_ui":"Selección múltiple","label":"Incidencia relacionada a delitos","name":"incidencia_delitos","required":False,
         "opciones":[
            "Disturbios en vía pública. (Riñas o Agresión)",
            "Daños a la propiedad. (Destruir, inutilizar o desaparecer).",
            "Extorsión (intimidar o amenazar a otras personas con fines de lucro).",
            "Hurto. (sustracción de artículos mediante el descuido).",
            "Receptación (persona que adquiere, recibe u oculta artículos provenientes de un delito en el que no participó).",
            "Contrabando (licor, cigarrillos, medicinas, ropa, calzado, etc.)",
            "Maltrato animal",
            "Tráfico ilegal de personas (coyotaje)",
            "Otro"
         ], "appearance":None, "choice_filter":None, "relevant":None},

        {"tipo_ui":"Selección múltiple","label":"Venta de drogas","name":"venta_drogas","required":False,
         "opciones":["bunker espacio cerrado","vía pública","exprés"], "appearance":None, "choice_filter":None, "relevant":None},

        {"tipo_ui":"Selección múltiple","label":"Delitos contra la vida","name":"delitos_vida","required":False,
         "opciones":["Homicidios","Heridos"], "appearance":None, "choice_filter":None, "relevant":None},

        {"tipo_ui":"Selección múltiple","label":"Delitos sexuales","name":"delitos_sexuales","required":False,
         "opciones":["Abuso sexual","Acoso sexual","Violación"], "appearance":None, "choice_filter":None, "relevant":None},

        {"tipo_ui":"Selección múltiple","label":"Asaltos","name":"asaltos","required":False,
         "opciones":["Asalto a personas","Asalto a comercio","Asalto a vivienda","Asalto a transporte público"],
         "appearance":None, "choice_filter":None, "relevant":None},

        {"tipo_ui":"Selección múltiple","label":"Estafas","name":"estafas","required":False,
         "opciones":["Billetes falso","Documentos falsos","Estafa (Oro)","Lotería falsos","Estafas informáticas","Estafa telefónica","Estafa con tarjetas"],
         "appearance":None, "choice_filter":None, "relevant":None},

        {"tipo_ui":"Selección múltiple","label":"Robo (Sustracción de artículos mediante la utilización de la fuerza)","name":"robo_fuerza","required":False,
         "opciones":["Tacha a comercio","Tacha a edificaciones","Tacha a vivienda","Tacha de vehículos",
                     "Robo de Ganado Abigeato (Destace de ganado)","Robo de bienes agrícola","Robo de vehículos","Robo de cable","Robo de combustible"],
         "appearance":None,"choice_filter":None,"relevant":None},

        {"tipo_ui":"Selección múltiple","label":"Explotación infantil","name":"explotacion_infantil","required":False,
         "opciones":["Sexual","Laboral"], "appearance":None, "choice_filter":None, "relevant":None},

        {"tipo_ui":"Selección múltiple","label":"Abandono de personas","name":"abandono_personas","required":False,
         "opciones":["Abandono de adulto mayor","Abandono de menor de edad","Abandono de incapaz"],
         "appearance":None, "choice_filter":None, "relevant":None},

        {"tipo_ui":"Selección múltiple","label":"Delitos Ambientales","name":"delitos_ambientales","required":False,
         "opciones":["Caza ilegal","Pesca ilegal","Tala ilegal"],
         "appearance":None, "choice_filter":None, "relevant":None},

        {"tipo_ui":"Selección múltiple","label":"Trata de personas","name":"trata_personas","required":False,
         "opciones":["Con fines laborales","Con fines sexuales"],
         "appearance":None, "choice_filter":None, "relevant":None},

        {"tipo_ui":"Selección múltiple","label":"Violencia Intrafamiliar (marque los que correspondan)","name":"vi_tipos","required":False,
         "opciones":[
            "Violencia psicológica (gritos, amenazas, burlas, maltratos, etc)",
            "Violencia física (golpes, empujones, etc)",
            "Violencia patrimonial (destrucción o retención de artículos, documentos, dinero, etc)",
            "Violencia sexual (actos sexuales no consentido)"
         ], "appearance":None, "choice_filter":None, "relevant":None},

        # ===================== PÁGINA 5: Riesgos sociales (NO obligatorias) =====================
        {"tipo_ui":"Selección múltiple","label":"Riesgos sociales","name":"riesgos_sociales","required":False,
         "opciones":[
            "Escándalos musicales.","Falta de oportunidades laborales.","Problemas Vecinales.",
            "Asentamientos ilegales (conocido como precarios).","Personas en situación de calle.",
            "Desvinculación escolar (deserción escolar)","Zona de prostitución","Consumo de alcohol en vía pública",
            "Personas con exceso de tiempo de ocio","Acumulación de basuras, aguas negras, mal alcantarillado.",
            "Carencia o inexistencia de alumbrado público.","Cuarterías","Lotes baldíos.","Ventas informales",
            "Pérdida de espacios públicos (lugares de esparcimiento como parques, polideportivos, etc.).","Otro"
         ], "appearance":None, "choice_filter":None, "relevant":None},

        {"tipo_ui":"Selección múltiple","label":"Falta de inversión social","name":"falta_inversion_social","required":False,
         "opciones":["Falta de oferta educativa","Falta de oferta deportiva","Falta de oferta recreativa","Falta de actividades culturales"],
         "appearance":None, "choice_filter":None, "relevant":None},

        {"tipo_ui":"Selección múltiple","label":"Consumo de drogas","name":"consumo_drogas","required":False,
         "opciones":["Área privada","Área pública"], "appearance":None, "choice_filter":None, "relevant":None},

        {"tipo_ui":"Selección múltiple","label":"Deficiencia en la infraestructura vial","name":"infra_vial","required":False,
         "opciones":["Calles en mal estado","Falta de señalización de tránsito","Carencia o inexistencia de aceras"],
         "appearance":None, "choice_filter":None, "relevant":None},

        {"tipo_ui":"Selección múltiple","label":"Búnker","name":"bunker","required":False,
         "opciones":["Casa de habitación","Edificación abandonada","Lote baldío","Otro"],
         "appearance":None, "choice_filter":None, "relevant":None},

        # ===================== PÁGINA 6: Información adicional =====================
        {"tipo_ui":"Selección única","label":"¿Usted tiene información de alguna persona o grupo que se dedique a realizar algún delito alrededor del local comercial? (confidencial)","name":"info_grupo_delito","required":True,
         "opciones":["Sí","No"], "appearance":None, "choice_filter":None, "relevant":None},

        {"tipo_ui":"Párrafo (texto largo)","label":"Si su respuesta es 'Sí', describa detalles (nombre del grupo, actividad ilícita, cantidad aproximada de miembros, personas/alias/señas físicas, domicilios/lugares, vehículos: placa, color y marca).","name":"detalle_grupo","required":True,
         "opciones":[], "appearance":None, "choice_filter":None,
         "relevant": f"${{info_grupo_delito}}='{slugify_name('Sí')}'"},

        {"tipo_ui":"Selección única","label":"¿Su local comercial ha sido víctima de algún delito en los últimos 12 meses? ¿Denunció ante el OIJ?","name":"victimizacion","required":True,
         "opciones":["NO he sido víctima de ningún delito","SI he sido víctima y SI denuncié","SI he sido víctima, pero NO denuncié."],
         "appearance":None, "choice_filter":None, "relevant":None},

        # Rama SI denunció
        {"tipo_ui":"Texto (corto)","label":"¿Cuál fue el delito del que fue víctima?","name":"delito_si","required":True,
         "opciones":[], "appearance":None, "choice_filter":None,
         "relevant": f"${{victimizacion}}='{slugify_name('SI he sido víctima y SI denuncié')}'"},

        {"tipo_ui":"Selección múltiple","label":"¿Cuál fue el modo de operar? (marque todos los factores pertinentes)","name":"modo_si","required":True,
         "opciones":["Arma blanca (cuchillo, machete, tijeras).","Arma de fuego.","Amenazas","Arrebato","Boquete","Ganzúa (pata de chancho)","Engaño","No sé.","Otro"],
         "appearance":None, "choice_filter":None,
         "relevant": f"${{victimizacion}}='{slugify_name('SI he sido víctima y SI denuncié')}'"},

        {"tipo_ui":"Selección única","label":"Horario del hecho delictivo","name":"horario_si","required":True,
         "opciones":["00:00 - 02:59 a. m.","03:00 - 05:59 a. m.","06:00 - 08:59 a. m.","09:00 - 11:59 a. m.","12:00 - 14:59 p. m.","15:00 - 17:59 p. m.","18:00 - 20:59 p. m.","21:00 - 23:59 p. m.","DESCONOCIDO"],
         "appearance":None, "choice_filter":None,
         "relevant": f"${{victimizacion}}='{slugify_name('SI he sido víctima y SI denuncié')}'"},

        # Rama NO denunció
        {"tipo_ui":"Texto (corto)","label":"¿Cuál fue el delito del que fue víctima?","name":"delito_no","required":True,
         "opciones":[], "appearance":None, "choice_filter":None,
         "relevant": f"${{victimizacion}}='{slugify_name('SI he sido víctima, pero NO denuncié.')}'"},

        {"tipo_ui":"Selección múltiple","label":"Motivo de NO denunciar","name":"motivo_no","required":True,
         "opciones":["Distancia (falta de oficinas para recepción de denuncias).","Miedo a represalias.","Falta de respuesta oportuna.",
                     "He realizado denuncias y no ha pasado nada.","Complejidad al colocar la denuncia.","Desconocimiento de dónde colocar la denuncia.",
                     "El Policía me dijo que era mejor no denunciar.","Falta de tiempo para colocar la denuncia"],
         "appearance":None, "choice_filter":None,
         "relevant": f"${{victimizacion}}='{slugify_name('SI he sido víctima, pero NO denuncié.')}'"},

        {"tipo_ui":"Selección múltiple","label":"¿Cuál fue el modo de operar? (marque todos los factores pertinentes)","name":"modo_no","required":True,
         "opciones":["Arma blanca (cuchillo, machete, tijeras).","Arma de fuego.","Amenazas","Arrebato","Boquete","Ganzúa (pata de chancho)","Engaño","No sé.","Otro"],
         "appearance":None, "choice_filter":None,
         "relevant": f"${{victimizacion}}='{slugify_name('SI he sido víctima, pero NO denuncié.')}'"},

        {"tipo_ui":"Selección única","label":"Horario del hecho delictivo","name":"horario_no","required":True,
         "opciones":["00:00 - 02:59 a. m.","03:00 - 05:59 a. m.","06:00 - 08:59 a. m.","09:00 - 11:59 a. m.","12:00 - 14:59 p. m.","15:00 - 17:59 p. m.","18:00 - 20:59 p. m.","21:00 - 23:59 p. m.","DESCONOCIDO"],
         "appearance":None, "choice_filter":None,
         "relevant": f"${{victimizacion}}='{slugify_name('SI he sido víctima, pero NO denuncié.')}'"},

        {"tipo_ui":"Selección única","label":"¿Cómo califica el servicio policial de la Fuerza Pública de Costa Rica alrededor de su local?","name":"fp_calificacion","required":True,
         "opciones":["Excelente","Bueno","Regular","Mala","Muy mala"], "appearance":None, "choice_filter":None, "relevant":None},

        {"tipo_ui":"Selección única","label":"¿Cómo ha sido el servicio de la Fuerza Pública en los últimos 24 meses?","name":"fp_24m","required":True,
         "opciones":["Mejor servicio","Igual","Peor servicio"], "appearance":None, "choice_filter":None, "relevant":None},

        {"tipo_ui":"Selección única","label":"¿Conoce usted a los policías de la Fuerza Pública de su zona comercial?","name":"conoce_policias","required":True,
         "opciones":["Sí","No"], "appearance":None, "choice_filter":None, "relevant":None},

        {"tipo_ui":"Selección única","label":"¿Conoce el programa de \"Seguridad Comercial\" que imparte Fuerza Pública?","name":"conoce_programa","required":True,
         "opciones":["Sí","No"], "appearance":None, "choice_filter":None, "relevant":None},

        {"tipo_ui":"Selección única","label":"¿Está inscrito en el programa de \"Seguridad Comercial\"?","name":"inscrito_programa","required":True,
         "opciones":["Sí","No"], "appearance":None, "choice_filter":None,
         "relevant": f"${{conoce_programa}}='{slugify_name('Sí')}'"},

        {"tipo_ui":"Selección única","label":"¿Le gustaría que se le contacte para formar parte del programa?","name":"desea_contacto","required":True,
         "opciones":["Sí","No"], "appearance":None, "choice_filter":None,
         "relevant": "(" + " or ".join([
            f"${{conoce_programa}}='{slugify_name('No')}'",
            f"${{inscrito_programa}}='{slugify_name('No')}'"
         ]) + ")"},

        {"tipo_ui":"Párrafo (texto largo)","label":"Si su respuesta es afirmativa, indicar nombre del comercio, correo electrónico y número de teléfono para contactarlo(a).","name":"datos_contacto","required":True,
         "opciones":[], "appearance":None, "choice_filter":None,
         "relevant": f"${{desea_contacto}}='{slugify_name('Sí')}'"},

        {"tipo_ui":"Párrafo (texto largo)","label":"¿Qué actividad considera que deba realizar la Fuerza Pública para mejorar la seguridad en su zona comercial? (opcional)","name":"actividad_fp","required":False,
         "opciones":[], "appearance":None, "choice_filter":None, "relevant":None},

        {"tipo_ui":"Párrafo (texto largo)","label":"¿Qué actividad considera que deba realizar la municipalidad para mejorar la seguridad en su zona comercial? (opcional)","name":"actividad_muni","required":False,
         "opciones":[], "appearance":None, "choice_filter":None, "relevant":None},

        {"tipo_ui":"Párrafo (texto largo)","label":"En el siguiente espacio podrá registrar alguna otra información que estime pertinente. (opcional)","name":"otra_info","required":False,
         "opciones":[], "appearance":None, "choice_filter":None, "relevant":None},

        {"tipo_ui":"Párrafo (texto largo)","label":"(Voluntario) Nombre, teléfono o correo electrónico para continuar colaborando de forma confidencial con Fuerza Pública. (opcional)","name":"contacto_voluntario","required":False,
         "opciones":[], "appearance":None, "choice_filter":None, "relevant":None},
    ]

    st.session_state.preguntas = seed
    st.session_state.seed_cargado = True
    st.session_state.perfil_seed = "comercio"
# ------------------------------------------------------------------------------------------
# Sidebar: Metadatos + Exportar/Importar proyecto
# ------------------------------------------------------------------------------------------
with st.sidebar:
    st.header("⚙️ Configuración")
    form_title = st.text_input("Título del formulario",
                               value=(f"Encuesta Comercio – {delegacion.strip()}" if delegacion.strip() else "Encuesta Comercio"))
    idioma = st.selectbox("Idioma por defecto (default_language)", options=["es","en"], index=0)
    version_auto = datetime.now().strftime("%Y%m%d%H%M")
    version = st.text_input("Versión (settings.version)", value=version_auto)

    st.markdown("---")
    st.caption("💾 Exporta/Importa tu proyecto (JSON)")
    col_exp, col_imp = st.columns(2)
    if col_exp.button("Exportar proyecto (JSON)", use_container_width=True):
        proj = {
            "form_title": form_title, "idioma": idioma, "version": version,
            "preguntas": st.session_state.preguntas,
            "reglas_visibilidad": st.session_state.reglas_visibilidad,
            "reglas_finalizar": st.session_state.reglas_finalizar
        }
        jbuf = BytesIO(json.dumps(proj, ensure_ascii=False, indent=2).encode("utf-8"))
        st.download_button("Descargar JSON", data=jbuf, file_name="proyecto_encuesta_comercio.json",
                           mime="application/json", use_container_width=True)
    up = col_imp.file_uploader("Importar JSON", type=["json"], label_visibility="collapsed")
    if up is not None:
        try:
            raw = up.read().decode("utf-8"); data = json.loads(raw)
            st.session_state.preguntas = list(data.get("preguntas", []))
            st.session_state.reglas_visibilidad = list(data.get("reglas_visibilidad", []))
            st.session_state.reglas_finalizar = list(data.get("reglas_finalizar", []))
            _rerun()
        except Exception as e:
            st.error(f"No se pudo importar el JSON: {e}")
# ------------------------------------------------------------------------------------------
# Constructor: Agregar nuevas preguntas
# ------------------------------------------------------------------------------------------
st.subheader("📝 Diseña tus preguntas")

with st.form("form_add_q", clear_on_submit=False):
    tipo_ui = st.selectbox("Tipo de pregunta", options=TIPOS)
    label = st.text_input("Etiqueta (texto exacto)")
    sugerido = slugify_name(label) if label else ""
    col_n1, col_n2, col_n3 = st.columns([2,1,1])
    name = col_n1.text_input("Nombre interno (XLSForm 'name')", value=sugerido)
    required = col_n2.checkbox("Requerida", value=False)
    appearance = col_n3.text_input("Appearance (opcional)", value="")
    opciones = []
    if tipo_ui in ("Selección única","Selección múltiple"):
        st.markdown("**Opciones (una por línea)**")
        txt_opts = st.text_area("Opciones", height=120)
        if txt_opts.strip(): opciones = [o.strip() for o in txt_opts.splitlines() if o.strip()]
    add = st.form_submit_button("➕ Agregar pregunta")

if add:
    if not label.strip():
        st.warning("Agrega una etiqueta.")
    else:
        base = slugify_name(name or label)
        usados = {q["name"] for q in st.session_state.preguntas}
        unico = asegurar_nombre_unico(base, usados)
        st.session_state.preguntas.append({
            "tipo_ui": tipo_ui, "label": label.strip(), "name": unico, "required": required,
            "opciones": opciones, "appearance": (appearance.strip() or None),
            "choice_filter": None, "relevant": None
        })
        st.success(f"Pregunta agregada: **{label}** (name: `{unico}`)")
# ------------------------------------------------------------------------------------------
# Lista / Ordenado / Edición (completa)
# ------------------------------------------------------------------------------------------
st.subheader("📚 Preguntas (ordénalas y edítalas)")
if not st.session_state.preguntas:
    st.info("Aún no has agregado preguntas.")
else:
    for idx, q in enumerate(st.session_state.preguntas):
        with st.container(border=True):
            c1, c2, c3, c4, c5 = st.columns([4, 2, 2, 2, 2])
            c1.markdown(f"**{idx+1}. {q['label']}**")
            meta = f"type: {q['tipo_ui']}  •  name: `{q['name']}`  •  requerida: {'sí' if q['required'] else 'no'}"
            if q.get("appearance"): meta += f"  •  appearance: `{q['appearance']}`"
            if q.get("choice_filter"): meta += f"  •  choice_filter: `{q['choice_filter']}`"
            if q.get("relevant"): meta += f"  •  relevant: `{q['relevant']}`"
            c1.caption(meta)
            if q["tipo_ui"] in ("Selección única","Selección múltiple"):
                c1.caption("Opciones: " + ", ".join(q.get("opciones") or []))

            up = c2.button("⬆️ Subir", key=f"up_{idx}", use_container_width=True, disabled=(idx == 0))
            down = c3.button("⬇️ Bajar", key=f"down_{idx}", use_container_width=True, disabled=(idx == len(st.session_state.preguntas)-1))
            edit = c4.button("✏️ Editar", key=f"edit_{idx}", use_container_width=True)
            borrar = c5.button("🗑️ Eliminar", key=f"del_{idx}", use_container_width=True)

            if up:
                st.session_state.preguntas[idx-1], st.session_state.preguntas[idx] = st.session_state.preguntas[idx], st.session_state.preguntas[idx-1]; _rerun()
            if down:
                st.session_state.preguntas[idx+1], st.session_state.preguntas[idx] = st.session_state.preguntas[idx], st.session_state.preguntas[idx+1]; _rerun()

            if edit:
                st.markdown("**Editar esta pregunta**")
                ne_label = st.text_input("Etiqueta", value=q["label"], key=f"e_label_{idx}")
                ne_name = st.text_input("Nombre interno (name)", value=q["name"], key=f"e_name_{idx}")
                ne_required = st.checkbox("Requerida", value=q["required"], key=f"e_req_{idx}")
                ne_appearance = st.text_input("Appearance", value=q.get("appearance") or "", key=f"e_app_{idx}")
                ne_choice_filter = st.text_input("choice_filter (opcional)", value=q.get("choice_filter") or "", key=f"e_cf_{idx}")
                ne_relevant = st.text_input("relevant (opcional – se autogenera por reglas)", value=q.get("relevant") or "", key=f"e_rel_{idx}")

                ne_opciones = q.get("opciones") or []
                if q["tipo_ui"] in ("Selección única","Selección múltiple"):
                    ne_opts_txt = st.text_area("Opciones (una por línea)", value="\n".join(ne_opciones), key=f"e_opts_{idx}")
                    ne_opciones = [o.strip() for o in ne_opts_txt.splitlines() if o.strip()]

                col_ok, col_cancel = st.columns(2)
                if col_ok.button("💾 Guardar cambios", key=f"e_save_{idx}", use_container_width=True):
                    new_base = slugify_name(ne_name or ne_label)
                    usados = {qq["name"] for j, qq in enumerate(st.session_state.preguntas) if j != idx}
                    ne_name_final = new_base if new_base not in usados else asegurar_nombre_unico(new_base, usados)

                    st.session_state.preguntas[idx]["label"] = ne_label.strip() or q["label"]
                    st.session_state.preguntas[idx]["name"] = ne_name_final
                    st.session_state.preguntas[idx]["required"] = ne_required
                    st.session_state.preguntas[idx]["appearance"] = ne_appearance.strip() or None
                    st.session_state.preguntas[idx]["choice_filter"] = ne_choice_filter.strip() or None
                    st.session_state.preguntas[idx]["relevant"] = ne_relevant.strip() or None
                    if q["tipo_ui"] in ("Selección única","Selección múltiple"):
                        st.session_state.preguntas[idx]["opciones"] = ne_opciones
                    st.success("Cambios guardados."); _rerun()
                if col_cancel.button("Cancelar", key=f"e_cancel_{idx}", use_container_width=True):
                    _rerun()

            if borrar:
                del st.session_state.preguntas[idx]
                st.warning("Pregunta eliminada."); _rerun()
# ------------------------------------------------------------------------------------------
# Condicionales
# ------------------------------------------------------------------------------------------
st.subheader("🔀 Condicionales (mostrar / finalizar)")
if not st.session_state.preguntas:
    st.info("Agrega preguntas para definir condicionales.")
else:
    with st.expander("👁️ Mostrar pregunta si se cumple condición", expanded=False):
        names = [q["name"] for q in st.session_state.preguntas]
        labels_by_name = {q["name"]: q["label"] for q in st.session_state.preguntas}

        target = st.selectbox("Pregunta a mostrar (target)", options=names, format_func=lambda n: f"{n} — {labels_by_name[n]}")
        src = st.selectbox("Depende de (source)", options=names, format_func=lambda n: f"{n} — {labels_by_name[n]}")
        op = st.selectbox("Operador", options=["=", "selected"])
        src_q = next((q for q in st.session_state.preguntas if q["name"] == src), None)

        vals = []
        if src_q and src_q["opciones"]:
            vals = st.multiselect("Valores (usa texto; internamente se usa slug)", options=src_q["opciones"])
            vals = [slugify_name(v) for v in vals]
        else:
            manual = st.text_input("Valor (si la pregunta no tiene opciones)")
            vals = [slugify_name(manual)] if manual.strip() else []

        if st.button("➕ Agregar regla de visibilidad"):
            if target == src:
                st.error("Target y Source no pueden ser la misma pregunta.")
            elif not vals:
                st.error("Indica al menos un valor.")
            else:
                st.session_state.reglas_visibilidad.append({"target": target, "src": src, "op": op, "values": vals})
                st.success("Regla agregada."); _rerun()

        if st.session_state.reglas_visibilidad:
            st.markdown("**Reglas de visibilidad actuales:**")
            for i, r in enumerate(st.session_state.reglas_visibilidad):
                st.write(f"- Mostrar **{r['target']}** si **{r['src']}** {r['op']} {r['values']}")
                if st.button(f"Eliminar regla #{i+1}", key=f"del_vis_{i}"):
                    del st.session_state.reglas_visibilidad[i]; _rerun()

    with st.expander("⏹️ Finalizar temprano si se cumple condición", expanded=False):
        names = [q["name"] for q in st.session_state.preguntas]
        labels_by_name = {q["name"]: q["label"] for q in st.session_state.preguntas}
        src2 = st.selectbox("Condición basada en", options=names, format_func=lambda n: f"{n} — {labels_by_name[n]}", key="final_src")
        op2 = st.selectbox("Operador", options=["=", "selected", "!="], key="final_op")
        src2_q = next((q for q in st.session_state.preguntas if q["name"] == src2), None)
        vals2 = []
        if src2_q and src2_q["opciones"]:
            vals2 = st.multiselect("Valores (slug interno)", options=src2_q["opciones"], key="final_vals")
            vals2 = [slugify_name(v) for v in vals2]
        else:
            manual2 = st.text_input("Valor (si no hay opciones)", key="final_manual")
            vals2 = [slugify_name(manual2)] if manual2.strip() else []
        if st.button("➕ Agregar regla de finalización"):
            if not vals2:
                st.error("Indica al menos un valor.")
            else:
                idx_src = next((i for i, q in enumerate(st.session_state.preguntas) if q["name"] == src2), 0)
                st.session_state.reglas_finalizar.append({"src": src2, "op": op2, "values": vals2, "index_src": idx_src})
                st.success("Regla agregada."); _rerun()

        if st.session_state.reglas_finalizar:
            st.markdown("**Reglas de finalización actuales:**")
            for i, r in enumerate(st.session_state.reglas_finalizar):
                st.write(f"- Si **{r['src']}** {r['op']} {r['values']} ⇒ ocultar lo que sigue (efecto fin)")
                if st.button(f"Eliminar regla fin #{i+1}", key=f"del_fin_{i}"):
                    del st.session_state.reglas_finalizar[i]; _rerun()
# ------------------------------------------------------------------------------------------
# Construcción XLSForm (incluye P1..P6) + constraints placeholders C/D/B
# ------------------------------------------------------------------------------------------
def _get_logo_media_name(): 
    return logo_media_name

def construir_xlsform(preguntas, form_title: str, idioma: str, version: str,
                      reglas_vis, reglas_fin):
    survey_rows = []; choices_rows = []

    vis_by_target = {}
    for r in reglas_vis:
        vis_by_target.setdefault(r["target"], []).append(
            {"src": r["src"], "op": r.get("op","="), "values": r.get("values",[])}
        )

    fin_conds = []
    for r in reglas_fin:
        cond = build_relevant_expr([{"src": r["src"], "op": r.get("op","="), "values": r.get("values",[])}])
        if cond: fin_conds.append((r["index_src"], cond))

    # Página 1: Intro (mismo encabezado con logo)
    survey_rows += [
        {"type":"begin_group","name":"p1_intro","label":"Introducción","appearance":"field-list"},
        {"type":"note","name":"intro_logo","label":form_title, "media::image": _get_logo_media_name()},
        {"type":"note","name":"intro_texto","label":INTRO_COMERCIO},
        {"type":"end_group","name":"p1_end"}
    ]

    # ======== SETS por página ========
    p2 = {"canton","distrito","barrio","edad","genero","escolaridad","tipo_local"}
    p3 = {"seguro_local","motivo_inseguro","comparativa_seguridad","motivo_comparativa"}
    p4 = {"incidencia_delitos","venta_drogas","delitos_vida","delitos_sexuales","asaltos","estafas",
          "robo_fuerza","explotacion_infantil","abandono_personas","delitos_ambientales","trata_personas","vi_tipos"}
    p5 = {"riesgos_sociales","falta_inversion_social","consumo_drogas","infra_vial","bunker"}
    p6 = {"info_grupo_delito","detalle_grupo","victimizacion",
          "delito_si","modo_si","horario_si",
          "delito_no","motivo_no","modo_no","horario_no",
          "fp_calificacion","fp_24m","conoce_policias","conoce_programa","inscrito_programa","desea_contacto","datos_contacto",
          "actividad_fp","actividad_muni","otra_info","contacto_voluntario"}

    def add_q(q, idx):
        x_type, default_app, list_name = map_tipo_to_xlsform(q["tipo_ui"], q["name"])

        rel_manual = q.get("relevant") or None
        rel_panel  = build_relevant_expr(vis_by_target.get(q["name"], []))
        nots = [xlsform_not(cond) for idx_src, cond in fin_conds if idx_src < idx]
        rel_fin = "(" + " and ".join(nots) + ")" if nots else None
        parts = [p for p in [rel_manual, rel_panel, rel_fin] if p]
        rel_final = parts[0] if parts and len(parts)==1 else ("(" + ") and (".join(parts) + ")" if parts else None)

        row = {"type": x_type, "name": q["name"], "label": q["label"]}
        if q.get("required"): row["required"] = "yes"
        app = q.get("appearance") or default_app
        if app: row["appearance"] = app
        if q.get("choice_filter"): row["choice_filter"] = q["choice_filter"]
        if rel_final: row["relevant"] = rel_final

        # Constraints placeholders C/D/B
        if q["name"] == "canton":
            row["constraint"] = ". != '__pick_canton__'"
            row["constraint_message"] = "Seleccione un cantón válido."
        if q["name"] == "distrito":
            row["constraint"] = ". != '__pick_distrito__'"
            row["constraint_message"] = "Seleccione un distrito válido."
        if q["name"] == "barrio":
            row["constraint"] = ". != '__pick_barrio__'"
            row["constraint_message"] = "Seleccione un barrio válido."

        survey_rows.append(row)

        # No generar opciones para C/D/B (se usan las del catálogo).
        if list_name and q["name"] not in {"canton","distrito","barrio"}:
            usados = set()
            for opt_label in (q.get("opciones") or []):
                base = slugify_name(opt_label)
                opt_name = asegurar_nombre_unico(base, usados); usados.add(opt_name)
                choices_rows.append({"list_name": list_name, "name": opt_name, "label": str(opt_label)})

    def add_page(group_name, page_label, names_set):
        survey_rows.append({"type":"begin_group","name":group_name,"label":page_label,"appearance":"field-list"})
        for i, q in enumerate(st.session_state.preguntas):
            if q["name"] in names_set: add_q(q, i)
        survey_rows.append({"type":"end_group","name":f"{group_name}_end"})

    add_page("p2_demograficos", "Datos demográficos", p2)
    add_page("p3_sentimiento", "Sentimiento de inseguridad comercial", p3)
    add_page("p4_incidencia", "Incidencia relacionada a delitos", p4)
    add_page("p5_riesgos", "Riesgos sociales", p5)
    add_page("p6_info_adicional", "Información adicional", p6)

    # Choices del catálogo manual (con unicidad por list+name)
    for r in st.session_state.choices_ext_rows:
        choices_rows.append(dict(r))

    # DataFrames
    survey_cols_all = set().union(*[r.keys() for r in survey_rows])
    survey_cols = [c for c in ["type","name","label","required","appearance","choice_filter","relevant","constraint","constraint_message","media::image"] if c in survey_cols_all]
    for k in sorted(survey_cols_all):
        if k not in survey_cols: survey_cols.append(k)
    df_survey = pd.DataFrame(survey_rows, columns=survey_cols)

    choices_cols_all = set()
    for r in choices_rows: choices_cols_all.update(r.keys())
    base_choice_cols = ["list_name","name","label"]
    for extra in sorted(choices_cols_all):
        if extra not in base_choice_cols: base_choice_cols.append(extra)
    df_choices = pd.DataFrame(choices_rows, columns=base_choice_cols) if choices_rows else pd.DataFrame(columns=base_choice_cols)

    df_settings = pd.DataFrame([{
        "form_title": form_title, "version": version, "default_language": idioma, "style": "pages"
    }], columns=["form_title","version","default_language","style"])

    return df_survey, df_choices, df_settings

def descargar_excel_xlsform(df_survey, df_choices, df_settings, nombre_archivo: str):
    buffer = BytesIO()
    with pd.ExcelWriter(buffer, engine="xlsxwriter") as writer:
        df_survey.to_excel(writer,  sheet_name="survey",   index=False)
        df_choices.to_excel(writer, sheet_name="choices",  index=False)
        df_settings.to_excel(writer, sheet_name="settings", index=False)
        wb = writer.book; fmt_hdr = wb.add_format({"bold": True, "align": "left"})
        for sheet, df in (("survey", df_survey), ("choices", df_choices), ("settings", df_settings)):
            ws = writer.sheets[sheet]; ws.freeze_panes(1, 0); ws.set_row(0, None, fmt_hdr)
            for col_idx, col_name in enumerate(list(df.columns)):
                ws.set_column(col_idx, col_idx, max(14, min(42, len(str(col_name)) + 8)))
    buffer.seek(0)
    st.download_button(
        label=f"📥 Descargar XLSForm ({nombre_archivo})",
        data=buffer, file_name=nombre_archivo,
        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        use_container_width=True
    )
# ------------------------------------------------------------------------------------------
# Exportar / Vista previa XLSForm
# ------------------------------------------------------------------------------------------
st.markdown("---")
st.subheader("📦 Generar XLSForm (Excel) para Survey123")
st.caption("""
Incluye:
- **survey** con `type,name,label,required,appearance,choice_filter,relevant,constraint,media::image`,
- **choices** (con `canton_key`, `distrito_key` y `any` para placeholders),
- **settings** con título, versión, idioma y **style = pages**.
""")

if st.button("🧮 Construir XLSForm", use_container_width=True, disabled=not st.session_state.preguntas):
    try:
        names = [q["name"] for q in st.session_state.preguntas]
        if len(names) != len(set(names)):
            st.error("Hay 'name' duplicados. Edita las preguntas para que cada 'name' sea único.")
        else:
            df_survey, df_choices, df_settings = construir_xlsform(
                st.session_state.preguntas,
                form_title=(f"Encuesta Comercio – {delegacion.strip()}" if delegacion.strip() else "Encuesta Comercio"),
                idioma=st.sidebar.session_state.get("idioma","es") if hasattr(st.sidebar, "session_state") else "es",
                version=(st.sidebar.session_state.get("version") if hasattr(st.sidebar, "session_state") else "") or datetime.now().strftime("%Y%m%d%H%M"),
                reglas_vis=st.session_state.reglas_visibilidad,
                reglas_fin=st.session_state.reglas_finalizar
            )
            st.success("XLSForm construido. Vista previa:")
            c1, c2, c3 = st.columns(3)
            c1.markdown("**Hoja: survey**");   c1.dataframe(df_survey,  use_container_width=True, hide_index=True)
            c2.markdown("**Hoja: choices**");  c2.dataframe(df_choices, use_container_width=True, hide_index=True)
            c3.markdown("**Hoja: settings**"); c3.dataframe(df_settings,use_container_width=True, hide_index=True)

            nombre_archivo = slugify_name(f"Encuesta Comercio – {delegacion.strip()}" if delegacion.strip() else "Encuesta Comercio") + "_xlsform.xlsx"
            descargar_excel_xlsform(df_survey, df_choices, df_settings, nombre_archivo)

            if st.session_state.get("_logo_bytes"):
                st.download_button("📥 Descargar logo para carpeta media",
                                   data=st.session_state["_logo_bytes"],
                                   file_name=logo_media_name, mime="image/png",
                                   use_container_width=True)

            st.info("Publica en Survey123 Connect: crea encuesta desde archivo, copia el logo a `media/` y publica.")
    except Exception as e:
        st.error(f"Ocurrió un error al generar el XLSForm: {e}")

# ------------------------------------------------------------------------------------------
# Exportar Word y PDF — helpers visuales
# ------------------------------------------------------------------------------------------
try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ROW_HEIGHT_RULE
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
except Exception:
    Document = None

try:
    from reportlab.pdfgen import canvas
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import cm
    from reportlab.lib.utils import ImageReader
    from reportlab.pdfbase.pdfmetrics import stringWidth
    from reportlab.lib.colors import HexColor, black
except Exception:
    canvas = None

def _build_cond_text(qname: str, reglas_vis: List[Dict]) -> str:
    rels = [r for r in reglas_vis if r.get("target") == qname]
    if not rels: return ""
    parts = []
    for r in rels:
        op = r.get("op", "="); vals = r.get("values", [])
        vtxt = ", ".join(vals) if vals else ""
        parts.append(f"{r['src']} {op} [{vtxt}]")
    return "Condición: se muestra si " + " OR ".join(parts)

def _get_logo_bytes_fallback() -> bytes | None:
    if st.session_state.get("_logo_bytes"): return st.session_state["_logo_bytes"]
    try:
        with open("001.png", "rb") as f: return f.read()
    except Exception:
        return None

def _wrap_text_lines(text: str, font_name: str, font_size: float, max_width: float) -> List[str]:
    if not text: return []
    from reportlab.pdfbase.pdfmetrics import stringWidth
    words = text.split(); lines, current = [], ""
    for w in words:
        test = (current + " " + w).strip()
        if stringWidth(test, font_name, font_size) <= max_width:
            current = test
        else:
            if current: lines.append(current)
            if stringWidth(w, font_name, font_size) > max_width:
                chunk = ""
                for ch in w:
                    if stringWidth(chunk + ch, font_name, font_size) <= max_width: chunk += ch
                    else:
                        if chunk: lines.append(chunk)
                        chunk = ch
                current = chunk
            else:
                current = w
    if current: lines.append(current)
    return lines

def _is_yes_no_options(opts: List[str]) -> bool:
    if not opts: return False
    norm = {slugify_name(x) for x in opts if x and str(x).strip()}
    yes_variants = {"si","sí","yes"}; no_variants = {"no"}
    return norm.issubset(yes_variants | no_variants) and any(y in norm for y in yes_variants) and any(n in norm for n in no_variants)

def _should_show_options(q: Dict) -> bool:
    if q.get("tipo_ui") not in ("Selección única", "Selección múltiple"): return False
    opts = q.get("opciones") or []
    return bool(opts) and not _is_yes_no_options(opts)
# ------------------------------------------------------------------------------------------
# DOCX: helpers de estilo y export
# ------------------------------------------------------------------------------------------
from docx import Document as _Doc
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ROW_HEIGHT_RULE

def _set_cell_shading(cell, fill_hex: str):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    shd = tcPr.find(qn('w:shd'))
    if shd is None:
        shd = OxmlElement('w:shd'); tcPr.append(shd)
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), fill_hex.replace('#','').upper())

def _set_cell_borders(cell, color_hex: str):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    borders = tcPr.find(qn('w:tcBorders'))
    if borders is None:
        borders = OxmlElement('w:tcBorders'); tcPr.append(borders)
    for edge in ('top','left','bottom','right'):
        tag = OxmlElement(f'w:{edge}'); tag.set(qn('w:val'), 'single'); tag.set(qn('w:sz'), '8'); tag.set(qn('w:color'), color_hex.replace('#','').upper())
        borders.append(tag)

def _add_observation_box(doc: _Doc, fill_hex: str, border_hex: str):
    from docx.shared import Inches
    tbl = doc.add_table(rows=1, cols=1); tbl.alignment = WD_TABLE_ALIGNMENT.LEFT; tbl.autofit = True
    cell = tbl.cell(0, 0); _set_cell_shading(cell, fill_hex); _set_cell_borders(cell, border_hex)
    row = tbl.rows[0]; row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST; row.height = Inches(1.1)
    cell.paragraphs[0].add_run("")

ALL_BY_PAGE_PRINT = [
    ("Introducción (P1)", {"__INTRO__"}),
    ("Sentimiento de inseguridad comercial (P3)", {"seguro_local","motivo_inseguro","comparativa_seguridad","motivo_comparativa"}),
    ("Incidencia relacionada a delitos (P4)", {"incidencia_delitos","venta_drogas","delitos_vida","delitos_sexuales","asaltos","estafas",
                                               "robo_fuerza","explotacion_infantil","abandono_personas","delitos_ambientales","trata_personas","vi_tipos"}),
    ("Riesgos sociales (P5)", {"riesgos_sociales","falta_inversion_social","consumo_drogas","infra_vial","bunker"}),
    ("Información adicional (P6)", {"info_grupo_delito","detalle_grupo","victimizacion","delito_si","modo_si","horario_si","delito_no","motivo_no",
                                    "modo_no","horario_no","fp_calificacion","fp_24m","conoce_policias","conoce_programa","inscrito_programa",
                                    "desea_contacto","datos_contacto","actividad_fp","actividad_muni","otra_info","contacto_voluntario"}),
]

def export_docx_form(preguntas: List[Dict], form_title: str, intro: str, reglas_vis: List[Dict]):
    if Document is None:
        st.error("Falta dependencia: instala `python-docx` para generar Word."); return
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    fills = ["#E6F4EA", "#E7F0FE", "#FDECEA"]; borders = ["#1E8E3E", "#1A73E8", "#D93025"]; BLACK = RGBColor(0,0,0)

    doc = Document()
    p = doc.add_paragraph(); run = p.add_run(form_title); run.bold = True; run.font.size = Pt(24); run.font.color.rgb = BLACK
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    logo_b = _get_logo_bytes_fallback()
    if logo_b:
        try:
            img_buf = BytesIO(logo_b)
            doc.add_paragraph().alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_picture(img_buf, width=Inches(2.8))
        except Exception:
            pass

    intro_p = doc.add_paragraph(intro); intro_p.runs[0].font.size = Pt(12); intro_p.runs[0].font.color.rgb = BLACK

    i = 1; color_idx = 0
    for section_title, names in ALL_BY_PAGE_PRINT:
        if "__INTRO__" in names: 
            continue
        sec = doc.add_paragraph(section_title)
        rs = sec.runs[0]; rs.bold = True; rs.font.size = Pt(14); rs.font.color.rgb = BLACK

        for q in preguntas:
            if q.get("name") not in names: continue
            doc.add_paragraph("")
            h = doc.add_paragraph(f"{i}. {q['label']}"); r = h.runs[0]; r.font.size = Pt(11); r.font.color.rgb = BLACK

            cond_txt = _build_cond_text(q["name"], reglas_vis)
            if cond_txt:
                cpara = doc.add_paragraph(cond_txt); rc = cpara.runs[0]; rc.italic = True; rc.font.size = Pt(9); rc.font.color.rgb = BLACK

            if _should_show_options(q):
                opts_str = ", ".join([str(x) for x in q.get("opciones") if str(x).strip()])
                opara = doc.add_paragraph(f"Opciones: {opts_str}"); ro = opara.runs[0]; ro.font.size = Pt(10); ro.font.color.rgb = BLACK

            fill = fills[color_idx % len(fills)]; border = borders[color_idx % len(borders)]; color_idx += 1
            _add_observation_box(doc, fill, border)
            help_p = doc.add_paragraph("Agregue sus observaciones sobre la pregunta.")
            rh = help_p.runs[0]; rh.italic = True; rh.font.size = Pt(9); rh.font.color.rgb = BLACK
            i += 1

    buf = BytesIO(); doc.save(buf); buf.seek(0)
    st.download_button("📄 Descargar Word del formulario", data=buf,
                       file_name=slugify_name(form_title) + "_formulario.docx",
                       mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                       use_container_width=True)

# ------------------------------------------------------------------------------------------
# PDF editable: export
# ------------------------------------------------------------------------------------------
def export_pdf_editable_form(preguntas: List[Dict], form_title: str, intro: str, reglas_vis: List[Dict]):
    if canvas is None:
        st.error("Falta dependencia: instala `reportlab` para generar PDF."); return
    from reportlab.pdfgen import canvas as _canvas
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import cm
    from reportlab.lib.utils import ImageReader
    from reportlab.lib.colors import HexColor, black

    PAGE_W, PAGE_H = A4; margin = 2 * cm; max_text_w = PAGE_W - 2 * margin
    title_font, title_size = "Helvetica-Bold", 24
    intro_font, intro_size = "Helvetica", 12; intro_line_h = 18
    sec_font, sec_size = "Helvetica-Bold", 14
    label_font, label_size = "Helvetica", 11
    cond_font, cond_size = "Helvetica-Oblique", 9
    helper_font, helper_size = "Helvetica-Oblique", 9
    opts_font, opts_size = "Helvetica", 10
    fills = [HexColor("#E6F4EA"), HexColor("#E7F0FE"), HexColor("#FDECEA")]
    borders = [HexColor("#1E8E3E"), HexColor("#1A73E8"), HexColor("#D93025")]
    field_h = 80; line_h = 14; y = PAGE_H - margin

    c = _canvas.Canvas(BytesIO(), pagesize=A4); buf = c._filename
    c = _canvas.Canvas(buf, pagesize=A4); c.setTitle(form_title)

    logo_b = _get_logo_bytes_fallback()
    if logo_b:
        try:
            img = ImageReader(BytesIO(logo_b))
            logo_w, logo_h = 160, 115
            c.drawImage(img, (PAGE_W - logo_w)/2, y - logo_h, width=logo_w, height=logo_h, preserveAspectRatio=True, mask='auto')
            y -= (logo_h + 24)
        except Exception:
            pass

    c.setFillColor(black); c.setFont(title_font, title_size); c.drawCentredString(PAGE_W/2, y, form_title); y -= 26
    c.setFont(intro_font, intro_size)
    for line in _wrap_text_lines(intro, intro_font, intro_size, max_text_w):
        if y < margin + 80: c.showPage(); y = PAGE_H - margin; c.setFillColor(black); c.setFont(intro_font, intro_size)
        c.drawString(margin, y, line); y -= intro_line_h

    c.showPage(); y = PAGE_H - margin; c.setFillColor(black)

    i = 1; color_idx = 0
    for section_title, names in ALL_BY_PAGE_PRINT:
        if "__INTRO__" in names: continue
        c.setFont(sec_font, sec_size); c.drawString(margin, y, section_title); y -= (line_h + 6); c.setFont(label_font, label_size)

        for q in st.session_state.preguntas:
            if q.get("name") not in names: continue

            label_lines = _wrap_text_lines(f"{i}. {q['label']}", label_font, label_size, max_text_w)
            needed = line_h * len(label_lines) + field_h + 26

            cond_txt = _build_cond_text(q["name"], reglas_vis)
            cond_lines = _wrap_text_lines(cond_txt, cond_font, cond_size, max_text_w) if cond_txt else []
            needed += line_h * len(cond_lines)

            opts_lines = []
            if _should_show_options(q):
                opts_str = ", ".join([str(x) for x in q.get("opciones") if str(x).strip()])
                opts_lines = _wrap_text_lines(f"Opciones: {opts_str}", opts_font, opts_size, max_text_w)
                needed += line_h * len(opts_lines)

            if y - needed < margin:
                c.showPage(); y = PAGE_H - margin; c.setFillColor(black)
                c.setFont(sec_font, sec_size); c.drawString(margin, y, section_title)
                y -= (line_h + 6); c.setFont(label_font, label_size)

            for line in label_lines: c.drawString(margin, y, line); y -= line_h
            if cond_lines:
                c.setFont(cond_font, cond_size)
                for cl in cond_lines: c.drawString(margin, y, cl); y -= line_h
                c.setFont(label_font, label_size)
            if opts_lines:
                c.setFont(opts_font, opts_size)
                for ol in opts_lines: c.drawString(margin, y, ol); y -= line_h
                c.setFont(label_font, label_size)

            fill_color = fills[color_idx % len(fills)]; border_color = borders[color_idx % len(borders)]; color_idx += 1
            c.setFillColor(fill_color); c.setStrokeColor(border_color)
            c.rect(margin, y - field_h, max_text_w, field_h, fill=1, stroke=1); c.setFillColor(black)
            c.acroForm.textfield(name=f"campo_obs_{i}", tooltip=f"Observaciones para: {q['name']}",
                                 x=margin, y=y - field_h, width=max_text_w, height=field_h,
                                 borderWidth=1, borderStyle='solid', forceBorder=True, fieldFlags=4096, value="")
            c.setFont(helper_font, helper_size); c.drawString(margin, y - field_h - 10, "Agregue sus observaciones sobre la pregunta.")
            c.setFont(label_font, label_size); y -= (field_h + 26); i += 1

        if y < margin + 120: c.showPage(); y = PAGE_H - margin; c.setFillColor(black)

    c.showPage(); c.save()
    pdf_buf = c._filename; data = pdf_buf.getvalue() if hasattr(pdf_buf,"getvalue") else pdf_buf
    st.download_button("🧾 Descargar PDF editable del formulario", data=data,
                       file_name=slugify_name(form_title) + "_formulario_editable.pdf",
                       mime="application/pdf", use_container_width=True)

# ---------- Botones Word/PDF ----------
st.markdown("### 📝 Exportar formulario en **Word** y **PDF editable**")
col_w, col_p = st.columns(2)
if col_w.button("Generar Word (DOCX)"):
    export_docx_form(st.session_state.preguntas,
                     form_title=(f"Encuesta Comercio – {delegacion.strip()}" if delegacion.strip() else "Encuesta Comercio"),
                     intro=INTRO_COMERCIO, reglas_vis=st.session_state.reglas_visibilidad)
if col_p.button("Generar PDF editable"):
    export_pdf_editable_form(st.session_state.preguntas,
                             form_title=(f"Encuesta Comercio – {delegacion.strip()}" if delegacion.strip() else "Encuesta Comercio"),
                             intro=INTRO_COMERCIO, reglas_vis=st.session_state.reglas_visibilidad)


